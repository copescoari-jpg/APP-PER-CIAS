#!/usr/bin/env python3
"""
SISTEMA ARI — Laudos Periciais e Respostas
Ari Vladimir Copesco Junior | CREA 060097553-3
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
import threading
import base64
import zipfile
import re
import json
import os
import subprocess
from datetime import date
from pathlib import Path
import anthropic
from docx import Document
from docx.shared import Pt, Emu, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io

# ───────────────────────────────────────────────────────────────────────────────
# CONFIGURAÇÕES
# ───────────────────────────────────────────────────────────────────────────────

CONFIG_FILE    = Path.home() / ".sistema_ari" / "config.json"
SISTEMA_ARI    = Path(r"C:\Users\ari\OneDrive\Documentos\SISTEMA ARI")
LAUDOS_PATH    = SISTEMA_ARI / "LAUDOS"
IMAGE_EXT      = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"}
MAX_PHOTOS     = 20
MAX_IMAGE_SIZE = (1280, 1280)
MODEL          = "claude-sonnet-4-6"

MESES_PT = {
    "January": "janeiro", "February": "fevereiro", "March": "março",
    "April": "abril",     "May": "maio",            "June": "junho",
    "July": "julho",      "August": "agosto",       "September": "setembro",
    "October": "outubro", "November": "novembro",   "December": "dezembro",
}

# ───────────────────────────────────────────────────────────────────────────────
# TEMA VISUAL — pensado para uso confortável e alto contraste
# ───────────────────────────────────────────────────────────────────────────────

COR_FUNDO       = "#eef1f5"   # fundo geral (azul-acinzentado claro)
COR_FUNDO_HDR   = "#1f3a5c"   # cabeçalho (azul-petróleo profundo)
COR_FUNDO_CARD  = "#ffffff"   # cartões dos passos
COR_DESTAQUE    = "#dde4ee"   # caixinha do número do passo
COR_TEXTO       = "#14171d"   # texto principal (alto contraste)
COR_TEXTO_FRACO = "#56607a"   # textos auxiliares / status
COR_ACAO        = "#1f3a5c"   # botões principais
COR_ACAO_HOVER  = "#162a44"
COR_SECUNDARIO  = "#3d6aa8"   # botões secundários
COR_SEC_HOVER   = "#27497f"
COR_OK          = "#1a6b3a"
COR_AVISO       = "#a85a00"
COR_ERRO        = "#a01a1a"
COR_BORDA       = "#c8d0db"
COR_OBS_BG      = "#fffef0"   # área de observações
COR_OBS_BORDA   = "#c8b96a"

FONTE_BASE      = "Segoe UI"
FONT_TITULO     = (FONTE_BASE, 24, "bold")
FONT_SUBTITULO  = (FONTE_BASE, 14)
FONT_PASSO      = (FONTE_BASE, 17, "bold")
FONT_LABEL      = (FONTE_BASE, 15)
FONT_ENTRY      = (FONTE_BASE, 15)
FONT_BTN        = (FONTE_BASE, 15, "bold")
FONT_BTN_GR     = (FONTE_BASE, 19, "bold")
FONT_AJUDA      = (FONTE_BASE, 13)
FONT_AJUDA_SM   = (FONTE_BASE, 12)
FONT_STATUS     = (FONTE_BASE, 14)
FONT_PASSO_NUM  = (FONTE_BASE, 22, "bold")

ALTURA_BTN      = 56
ALTURA_BTN_GR   = 80
ALTURA_ENTRY    = 48

# ───────────────────────────────────────────────────────────────────────────────
# UTILITÁRIOS
# ───────────────────────────────────────────────────────────────────────────────

def load_config() -> dict:
    if CONFIG_FILE.exists():
        try:
            return json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {"api_key": ""}

def save_config(cfg: dict):
    CONFIG_FILE.parent.mkdir(parents=True, exist_ok=True)
    CONFIG_FILE.write_text(json.dumps(cfg, indent=2, ensure_ascii=False), encoding="utf-8")

def extract_docx_text(path: str) -> str:
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.open("word/document.xml").read().decode("utf-8")
        text = re.sub(r"<[^>]+>", " ", xml)
        return re.sub(r"\s+", " ", text).strip()
    except Exception as e:
        return f"[Erro ao ler .docx: {e}]"

def _extract_docx_lines(path: str) -> str:
    """Extrai texto do .docx linha a linha via python-docx (preserva estrutura)."""
    try:
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return "\n".join(lines)
    except Exception:
        return ""

def _extrair_partes_do_nome(nome: str) -> dict:
    """Tenta extrair 'Reclamante x Reclamada' de um nome de arquivo ou pasta."""
    # Remove prefixos comuns de documentos
    clean = re.sub(
        r'(?i)^\s*(?:pr[eé][\-_\s]*laudo|pre[\-_]?laudo|laudo|campo|'
        r'anota\w*|dilig\w*|secretar\w*|processo)\s*[\-_]?\s*',
        '', nome
    ).strip()
    # Remove sufixos do nosso próprio padrão de nome: "= Insalubr. = Função"
    # para que a reclamada não inclua esses campos adicionais
    clean = re.split(r'\s+=\s+', clean)[0].strip()
    m = re.match(r'^(.+?)\s+[xX]\s+(.+)$', clean)
    if m:
        return {'reclamante': m.group(1).strip(), 'reclamada': m.group(2).strip()}
    return {}


def _extrair_campo(texto: str, label: str) -> str:
    """
    Extrai o valor de um campo em múltiplos formatos:
      - 'Label: Valor'  ou  'Label - Valor'
      - 'Label | Valor'  (células de tabela)
      - 'Label\\nValor'  (label numa linha, valor na seguinte)
    """
    linhas = texto.split('\n')
    for i, linha in enumerate(linhas):
        # Mesmo linha: Label: Valor  /  Label - Valor  /  Label | Valor
        m = re.match(rf'(?i)\s*{label}\s*[:\-\|]\s*(.+)', linha)
        if m:
            val = re.split(r'\s{3,}|\t|\|', m.group(1))[0].strip()
            val = re.sub(r'\s+', ' ', val).strip()
            if 3 <= len(val) <= 120:
                return val

        # Label sozinho na linha → valor na próxima linha não vazia
        if re.match(rf'(?i)^\s*{label}\s*$', linha.strip()):
            for j in range(i + 1, min(i + 3, len(linhas))):
                val = linhas[j].strip()
                if val:
                    val = re.split(r'\s{3,}|\t|\|', val)[0].strip()
                    val = re.sub(r'\s+', ' ', val).strip()
                    if 3 <= len(val) <= 120:
                        return val
                    break
    return ''


def parse_processo(path: str) -> dict:
    """
    Extrai Reclamante, Reclamada, Função e tipo de um arquivo (.docx/.pdf).
    Tenta: (1) nome do arquivo, (2) conteúdo do documento.
    """
    data = {}

    # ── 1. Nome do arquivo ─────────────────────────────────────────────────
    partes = _extrair_partes_do_nome(Path(path).stem)
    data.update(partes)

    # ── 2. Conteúdo do documento ───────────────────────────────────────────
    ext = Path(path).suffix.lower()
    if ext == '.docx':
        text = _extract_docx_lines(path)
    elif ext == '.pdf':
        text = extract_pdf_text(path)
    else:
        return data

    if not text:
        return data

    for label, key in [('Reclamante', 'reclamante'), ('Reclamada', 'reclamada')]:
        if key not in data:
            val = _extrair_campo(text, label)
            if val:
                data[key] = val

    if 'funcao' not in data:
        for label in [r'Fun[çc][ãa]o', 'Cargo', r'Fun[çc][ãa]o\s*/\s*Cargo']:
            val = _extrair_campo(text, label)
            if val:
                data['funcao'] = val
                break

    data['insalubr']  = bool(re.search(r'(?i)insalubr',    text))
    data['periculos'] = bool(re.search(r'(?i)periculosid', text))

    return data


def parse_processo_da_pasta(folder: str, det: dict) -> dict:
    """
    Extrai dados do processo combinando: nome da pasta → pré-laudo → campo.
    det é o resultado de auto_detect().
    """
    data = {}

    # 1. Nome da pasta do processo
    partes = _extrair_partes_do_nome(Path(folder).name)
    data.update(partes)

    # 2. Pré-laudo (arquivo mais confiável para dados estruturados)
    if det.get('pre_laudo') and not (data.get('reclamante') and data.get('reclamada')):
        try:
            d = parse_processo(det['pre_laudo'])
            for k, v in d.items():
                if k not in data or not data[k]:
                    data[k] = v
        except Exception:
            pass

    # 3. Campo (complementar)
    if det.get('campo') and not (data.get('reclamante') and data.get('reclamada')):
        try:
            d = parse_processo(det['campo'])
            for k, v in d.items():
                if k not in data or not data[k]:
                    data[k] = v
        except Exception:
            pass

    return data

def extract_pdf_text(path: str) -> str:
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n".join(parts)
    except Exception as e:
        return f"[Erro ao ler .pdf: {e}]"

def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        return extract_docx_text(path)
    if ext == ".pdf":
        return extract_pdf_text(path)
    return "[Formato não suportado]"

def encode_image(path: str) -> tuple:
    with Image.open(path) as img:
        img.thumbnail(MAX_IMAGE_SIZE, Image.Resampling.LANCZOS)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
    return base64.standard_b64encode(buf.getvalue()).decode(), "image/jpeg"

def get_photos(folder: str) -> list:
    return sorted(
        [f for f in Path(folder).iterdir() if f.suffix.lower() in IMAGE_EXT]
    )[:MAX_PHOTOS]

def today_pt() -> str:
    d = date.today().strftime("%d de %B de %Y")
    for en, pt in MESES_PT.items():
        d = d.replace(en, pt)
    return d

def load_profile_files() -> tuple[str, str]:
    agente = (SISTEMA_ARI / "AGENTE_HUMANIZADOR.md")
    perfil  = (SISTEMA_ARI / "PERFIL_ESCRITA_ARI.md")
    return (
        agente.read_text(encoding="utf-8") if agente.exists() else "",
        perfil.read_text(encoding="utf-8")  if perfil.exists()  else "",
    )

def load_reference_laudo() -> str:
    """Carrega o primeiro laudo real como exemplo de formato."""
    if LAUDOS_PATH.exists():
        for f in sorted(LAUDOS_PATH.glob("*.docx")):
            txt = extract_docx_text(str(f))
            if len(txt) > 800:
                return txt[:7000]
    return ""

AVALIACAO_KEYWORDS = [
    "ruido", "ruído",
    "calor", "ibutg", "temperatur",
    "vibrac",                    # vibração / vibracao
    "rni",                       # radiação não ionizante
    "ltcat", "ppra", "pcmso", "pgr",
    "dosim",                     # dosimetria
    "medic", "mediç",            # medição
    "aval-",                     # "Avaliação-..." (sem casar com 'avaliacao' do pré-laudo)
    "avaliacao_", "avaliação_",
    "nho-",                      # NHO-01, NHO-06, NHO-09
    "agente",                    # "agentes_quimicos", "agente_fisico"
    "iluminanc",                 # iluminância
    "quimic",                    # químicos
]
DOCX_PDF = {".docx", ".pdf"}

def _eh_avaliacao(nome_lower: str) -> bool:
    return any(k in nome_lower for k in AVALIACAO_KEYWORDS)


def auto_detect(folder: str) -> dict:
    fp = Path(folder)
    r = {"pre_laudo": None, "campo": None,
         "photos_sub": None, "photo_count": 0,
         "laudo": None, "avaliacoes": []}
    for item in fp.iterdir():
        nl = item.name.lower()
        if item.is_dir():
            fotos = [f for f in item.iterdir()
                     if f.is_file() and f.suffix.lower() in IMAGE_EXT]
            if len(fotos) > r["photo_count"]:
                r["photos_sub"]  = str(item)
                r["photo_count"] = len(fotos)
        elif item.is_file() and item.suffix.lower() in DOCX_PDF:
            if any(k in nl for k in ["pre-laudo","pre_laudo","prelaudo","preliminar","secretari"]):
                r["pre_laudo"] = str(item)
            elif any(k in nl for k in ["campo","anotac","diligencia"]):
                r["campo"] = str(item)
            elif nl.startswith("laudo") and "impugnac" not in nl and "esclar" not in nl:
                r["laudo"] = str(item)
            elif _eh_avaliacao(nl):
                r["avaliacoes"].append(item)
    r["avaliacoes"].sort(key=lambda p: p.name.lower())
    return r

def output_path(folder: str, prefix: str) -> str:
    # Limita o nome do arquivo para que o caminho total fique abaixo de 259 chars.
    # Windows MAX_PATH = 260; reservamos: pasta + sep(1) + .docx(5) = pasta + 6
    folder_len = len(str(Path(folder)))
    max_nome = max(20, 259 - folder_len - 6)
    prefix = prefix[:max_nome]

    p = Path(folder) / f"{prefix}.docx"
    n = 2
    while p.exists():
        sfx = f"_{n}"
        p = Path(folder) / f"{prefix[:max_nome - len(sfx)]}{sfx}.docx"
        n += 1
    return str(p)

def _sanitize_folder_name(name: str) -> str:
    """Remove caracteres inválidos em nomes de pasta no Windows."""
    name = re.sub(r'[<>:"/\\|?*\n\r\t]', '', name).strip()
    name = re.sub(r'\s+', ' ', name)
    return name[:120].rstrip('. ')


def pasta_destino_laudo(processo_folder: str, reclamante: str, reclamada: str) -> Path:
    """Retorna (e cria) a subpasta '{Reclamante} x {Reclamada}' dentro do processo.

    Se faltarem nomes, devolve a própria pasta do processo (fallback seguro).
    """
    base = Path(processo_folder)
    rec = (reclamante or '').strip()
    red = (reclamada or '').strip()
    if not rec and not red:
        return base
    if rec and red:
        sub = f"{rec} x {red}"
    else:
        sub = rec or red
    destino = base / _sanitize_folder_name(sub)
    destino.mkdir(parents=True, exist_ok=True)
    return destino


def abrir_pasta_no_explorer(path: str | Path) -> None:
    """Abre o Windows Explorer na pasta indicada."""
    try:
        os.startfile(str(path))  # type: ignore[attr-defined]
    except Exception:
        try:
            subprocess.Popen(["explorer", str(path)])
        except Exception:
            pass


def build_laudo_filename(reclamante: str, reclamada: str,
                          insalubr: bool, periculos: bool, funcao: str) -> str:
    """Constrói nome do arquivo: Reclamante x Reclamada = Tipo = Função"""
    rec = reclamante.strip()
    red = reclamada.strip()
    fun = funcao.strip()

    if rec and red:
        base = f"{rec} x {red}"
    elif rec:
        base = rec
    elif red:
        base = red
    else:
        return "LAUDO_PERICIAL"

    parts = [base]
    if insalubr:
        parts.append("Insalubr.")
    if periculos:
        parts.append("Periculos.")
    if fun:
        parts.append(fun)

    name = " = ".join(parts)
    name = re.sub(r'[<>:"/\\|?*\n\r\t]', '', name).strip()
    return name[:180] or "LAUDO_PERICIAL"

# ───────────────────────────────────────────────────────────────────────────────
# SALVAR .DOCX  (com cabeçalho e rodapé)
# ───────────────────────────────────────────────────────────────────────────────

SECTION_HEADERS = [
    "OBJETIVO", "PROCESSO", "METODOLOGIA", "DILIGÊNCIAS",
    "ATIVIDADES DA RECLAMANTE", "ATIVIDADES DO RECLAMANTE",
    "DESCRIÇÃO DO LOCAL", "EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL",
    "RESULTADOS DAS AVALIAÇÕES", "RESPOSTAS AOS QUESITOS",
    "HONORÁRIOS PERICIAIS", "CONCLUSÃO", "ENCERRAMENTO",
    "Quesitos da Reclamante", "Quesitos da Reclamada",
    "Quesitos do Juízo", "Quesitos Complementares",
    "ESCLARECIMENTOS", "CONSIDERAÇÕES",
]
NUMBERED = {"6", "7", "8", "9", "10", "11", "12"}

def _is_header(line: str) -> bool:
    u = line.strip().upper()
    for n in NUMBERED:
        if u.startswith(f"{n} -") or u.startswith(f"{n}-"):
            return True
    return any(u == h.upper() or u.startswith(h.upper()) for h in SECTION_HEADERS)

def _add_docx_header_footer(sec):
    """Adiciona cabeçalho (LAUDO TÉCNICO + linha) e rodapé (endereço)."""
    # ── Cabeçalho ───────────────────────────────────────────────────────────
    hdr = sec.header

    # Parágrafo 0 — vazio, centralizado
    p0 = hdr.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parágrafo 1 — "LAUDO TÉCNICO" em cinza 26pt
    p1 = hdr.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("LAUDO TÉCNICO")
    run1.font.name = "Arial"
    run1.font.size = Pt(26)
    run1.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Parágrafo 2 — linha horizontal (borda inferior do parágrafo)
    p2 = hdr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p2._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '32')    # 4pt — equivalente ao traço do modelo
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ── Rodapé ──────────────────────────────────────────────────────────────
    ftr = sec.footer

    footer_lines = [
        "Ari Vladimir Copesco Júnior",
        "Engenheiro Segurança do Trabalho CREA 060097553-3",
        "Rua Marechal Rondon, n° 224 Fone (016) 3235-6763",
        "Ribeirão Preto/SP CEP 14.025-430",
    ]

    for i, line in enumerate(footer_lines):
        p = ftr.paragraphs[0] if i == 0 else ftr.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.italic = True

    # Parágrafo final vazio
    ftr.add_paragraph()


def save_docx(text: str, out: str):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin      = Emu(540385)
        sec.bottom_margin   = Emu(360045)
        sec.left_margin     = Emu(1080135)
        sec.right_margin    = Emu(900430)
        sec.header_distance = Emu(180340)
        sec.footer_distance = Emu(360045)
        _add_docx_header_footer(sec)

    sty = doc.styles["Normal"]
    sty.font.name = "Arial"
    sty.font.size = Pt(12)

    for raw in text.split("\n"):
        s = raw.strip()
        if not s:
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        if _is_header(s):
            p.paragraph_format.space_before = Pt(14)
            r = p.add_run(s)
            r.bold = True
        else:
            r = p.add_run(s)
        r.font.name = "Arial"
        r.font.size = Pt(12)

    doc.save(out)

# ───────────────────────────────────────────────────────────────────────────────
# API CLAUDE — LAUDO PERICIAL
# ───────────────────────────────────────────────────────────────────────────────

_REGRAS_LAUDO = """\
REGRAS ABSOLUTAS — LAUDO PERICIAL (prioridade máxima):

1. ESTRUTURA — 12 SEÇÕES NESTA ORDEM EXATA (sem pular, sem renomear):
   OBJETIVO → PROCESSO → METODOLOGIA → DILIGÊNCIAS → ATIVIDADES DA RECLAMANTE
   → 6 - DESCRIÇÃO DO LOCAL → 7 - EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL
   → 8 - RESULTADOS DAS AVALIAÇÕES → 9 - RESPOSTAS AOS QUESITOS
   → 10 - HONORÁRIOS PERICIAIS → 11 - CONCLUSÃO → 12 - ENCERRAMENTO

2. VOZ: SEMPRE plural majestático.
   ✅ "nossa análise", "concluímos", "verificamos", "optamos", "entendemos"
   ❌ PROIBIDO: "eu", "minha análise", "concluo"

3. PARTES: "Reclamante" e "Reclamada" SEMPRE com maiúscula inicial.

4. NORMAS: SEMPRE com Anexo e Portaria de origem.
   ✅ "Anexo nº 1 da NR-15 da Portaria 3.214/78"
   ❌ PROIBIDO: "conforme NR-15" sem especificar Anexo e Portaria

5. SEÇÃO 9 — cada resposta começa com "Resposta:"
   Negativa: "Resposta: não." ou "Resposta: Segundo nossa análise, não."
   Afirmativa: "Resposta: Sim." ou "Resposta: Segundo nossa análise, sim."
   Prejudicado: "Resposta: Prejudicado, tendo em vista a resposta ao quesito anterior."

6. FÓRMULAS FIXAS — copiar exatamente:

   OBJETIVO: "Este laudo visa descrever as condições e o processo de trabalho do Reclamante, com o objetivo de se detectar a insalubridade pleiteada e, se houver, enquadrá-la de acordo com a Norma Regulamentadora nº 15 e seus anexos da Portaria nº 3214/78."

   DILIGÊNCIAS (abertura): "Para atingir a correta interpretação dos fatos e encaminhamento final deste laudo pericial, sem subjetivismo e com embasamento técnico – legal, foi avaliada as condições em que trabalhava a Reclamante e feita uma minuciosa vistoria em seu local de trabalho. Esta atividade foi desenvolvida no dia [data]."

   SEÇÃO 8 (abertura obrigatória): "A palavra 'insalubre' vem do latim e significa tudo aquilo que origina doença, sendo que a insalubridade é a qualidade de insalubre. Já o conceito legal de insalubridade é dado pelo artigo 189 da Consolidação das Leis do Trabalho, nos seguintes termos: 'Serão consideradas atividades ou operações insalubres aquelas que, por sua natureza, condições ou métodos de trabalho, exponham os empregados a agentes nocivos à saúde, acima dos limites de tolerância, fixados em razão da natureza e da intensidade do agente e do tipo de exposição aos seus efeitos'."

   HONORÁRIOS: "Após ter cumprido a tarefa que lhe foi confiada, vem o perito requerer, com a devida cautela e respeito, o arbitramento de seus honorários periciais definidos em R$ 4.500,00 (Quatro mil e quinhentos reais)."

   CONCLUSÃO (abertura): "Pelo resultado das avaliações apresentadas no laudo pericial, onde foram relatados os riscos potenciais à saúde, sob o ponto de vista de Higiene e Segurança do Trabalho, e com embasamento nas declarações do citado documento, concluímos que:"

   ENCERRAMENTO: "Nada mais havendo a considerar, encerramos aqui os presentes trabalhos Periciais, compostos de 10 páginas com exceção desta que segue devidamente datada e assinada digitalmente através do sistema PJe-JT."
   "Ribeirão Preto, [data]."
   "Ari Vladimir Copesco Júnior"
   "Engº de Segurança do Trabalho"
   "CREA 060097553-3"

7. FORMATO: texto puro. SEM markdown (sem **, sem #, sem *, sem -). Parágrafos por \\n.
   SEM bullet points no corpo do laudo. Seções separadas por linha em branco.
"""

def gerar_laudo(api_key, pre_laudo, campo, photos, agente_md, perfil_md,
                ref_laudo="", obs="", avaliacoes="", progress_cb=None) -> str:
    client = anthropic.Anthropic(api_key=api_key)

    ref_block = (
        f"\nEXEMPLO REAL DE LAUDO DO ARI (referência de formato, vocabulário e estrutura):\n"
        f"{ref_laudo}\n(— fim do exemplo —)\n"
    ) if ref_laudo else ""

    obs_block = (
        f"\n=== OBSERVAÇÕES DO PERITO (considerar obrigatoriamente na elaboração) ===\n{obs}\n"
    ) if obs else ""

    aval_block = (
        f"\n=== AVALIAÇÕES TÉCNICAS REALIZADAS NA PERÍCIA ===\n{avaliacoes}\n"
    ) if avaliacoes else ""

    system_prompt = (
        "Você é o assistente de escrita de Ari Vladimir Copesco Junior,\n"
        "Engenheiro de Segurança do Trabalho e Perito Judicial — 3ª Vara do Trabalho\n"
        "de Ribeirão Preto — TRT 15ª Região. CREA 060097553-3.\n\n"
        + _REGRAS_LAUDO
        + "\n\nAGENTE HUMANIZADOR:\n" + agente_md
        + "\n\nPERFIL DE ESCRITA DETALHADO:\n" + perfil_md
        + ref_block
        + "\n\nRetorne APENAS o texto do laudo, sem comentários adicionais."
    )

    user_parts = []
    intro = (
        f"Elabore o laudo pericial COMPLETO com base nos documentos abaixo."
        f"{obs_block}\n"
        f"=== PRÉ-LAUDO (preparado pelas secretárias) ===\n{pre_laudo}\n\n"
        f"=== ANOTAÇÕES DE CAMPO / DADOS DA DILIGÊNCIA ===\n{campo}\n\n"
        f"{aval_block}"
        f"=== FOTOS DA DILIGÊNCIA ===\n"
    )
    for ph in photos:
        nome = ph.stem.replace("_", " ").replace("-", " ")
        intro += f"• {ph.name} — {nome}\n"
    intro += "\nImagens anexadas abaixo. Elabore o laudo COMPLETO agora."
    user_parts.append({"type": "text", "text": intro})

    for i, ph in enumerate(photos):
        if progress_cb:
            progress_cb(f"Carregando foto {i+1}/{len(photos)}: {ph.name}")
        try:
            data, mt = encode_image(str(ph))
            user_parts.append({"type": "image",
                                "source": {"type": "base64", "media_type": mt, "data": data}})
            user_parts.append({"type": "text",
                                "text": f"[Foto: {ph.stem.replace('_',' ').replace('-',' ')}]"})
        except Exception as e:
            user_parts.append({"type": "text", "text": f"[Foto {ph.name} — erro: {e}]"})

    user_parts.append({"type": "text",
                        "text": "Elabore agora o laudo pericial COMPLETO nas 12 seções."})

    if progress_cb:
        progress_cb("Gerando laudo com IA — aguarde, pode levar alguns minutos...")

    resp = client.messages.create(
        model=MODEL, max_tokens=16000,
        system=system_prompt,
        messages=[{"role": "user", "content": user_parts}],
    )
    return resp.content[0].text

# ───────────────────────────────────────────────────────────────────────────────
# API CLAUDE — RESPONDER IMPUGNAÇÃO / QUESITOS
# ───────────────────────────────────────────────────────────────────────────────

_REGRAS_RESPOSTA = """\
REGRAS ABSOLUTAS — ESCLARECIMENTOS DO PERITO (prioridade máxima):

ESTRUTURA OBRIGATÓRIA:
1. "EXMO(A) SR.(A) DR.(A) JUIZ(ÍZA) DA [X] VARA DO TRABALHO DE [CIDADE]/SP."
2. "Processo n°: [extrair do documento]"
3. "Reclamante: [extrair do documento]"
4. "Reclamada: [extrair do documento]"
5. Abertura: "Ari Vladimir Copesco Júnior, Engenheiro de Segurança do Trabalho, CREA 060097553-3, Perito Judicial designado nos autos da RECLAMAÇÃO TRABALHISTA em epígrafe, vem respeitosamente à presença de V. Exa. apresentar seus esclarecimentos em resposta às considerações apresentadas pelo(a) Assistente Técnico(a)."
6. Resposta técnica ponto a ponto a cada impugnação ou quesito, citando NRs, NHOs, ACGIH
7. Manutenção ou retificação das conclusões com justificativa técnica
8. "Nestes termos, aguarda deferimento."
9. "Ribeirão Preto, [data]."
10. "Ari Vladimir Copesco Júnior | Engº de Segurança do Trabalho | CREA 060097553-3 | Perito Judicial"

REGRAS DE VOZ E ESTILO:
- SEMPRE plural majestático: "nossa análise", "concluímos", "verificamos", "optamos"
- PROIBIDO: "eu", "minha"
- Imparcialidade técnica: se o questionamento tiver mérito, admita e esclareça
- Citar normas com Anexo e Portaria: "Anexo nº X da NR-15 da Portaria 3.214/78"
- Texto puro, SEM markdown, parágrafos separados por \\n
"""

def gerar_resposta(api_key, doc_recebido, meu_laudo, agente_md, perfil_md,
                   obs="", progress_cb=None) -> str:
    client = anthropic.Anthropic(api_key=api_key)

    laudo_ctx = (
        f"\nMEU LAUDO ORIGINAL (para referência e defesa das conclusões):\n{meu_laudo}\n"
    ) if meu_laudo else ""

    obs_block = (
        f"\n=== OBSERVAÇÕES DO PERITO (considerar obrigatoriamente) ===\n{obs}\n"
    ) if obs else ""

    system_prompt = (
        "Você é Ari Vladimir Copesco Júnior, Engenheiro de Segurança do Trabalho\n"
        "(CREA 060097553-3), Perito Judicial — 3ª Vara do Trabalho de Ribeirão Preto\n"
        "— TRT 15ª Região.\n\n"
        "Você recebeu impugnações ao seu laudo pericial ou quesitos complementares\n"
        "formulados pelo(a) Assistente Técnico(a) da parte.\n\n"
        + _REGRAS_RESPOSTA
        + "\n\nAGENTE HUMANIZADOR:\n" + agente_md
        + "\n\nPERFIL DE ESCRITA:\n" + perfil_md
        + f"\n\nData de hoje: {today_pt()}"
        + "\n\nRetorne APENAS o texto do documento, sem comentários adicionais."
    )

    user_content = (
        f"Elabore os esclarecimentos em resposta ao documento abaixo."
        f"{obs_block}\n"
        f"=== DOCUMENTO RECEBIDO (Impugnação ou Quesitos Complementares) ===\n"
        f"{doc_recebido}"
        f"{laudo_ctx}"
        f"\nElabore agora a resposta completa."
    )

    if progress_cb:
        progress_cb("Gerando resposta com IA — aguarde...")

    resp = client.messages.create(
        model=MODEL, max_tokens=8000,
        system=system_prompt,
        messages=[{"role": "user", "content": user_content}],
    )
    return resp.content[0].text

# ───────────────────────────────────────────────────────────────────────────────
# DIALOG — CONFIGURAÇÃO DA API
# ───────────────────────────────────────────────────────────────────────────────

class SettingsDialog(ctk.CTkToplevel):
    def __init__(self, parent, current_key: str, on_save):
        super().__init__(parent)
        self.title("Configuração da API Anthropic")
        self.geometry("640x280")
        self.resizable(False, False)
        self.grab_set()
        self.lift()
        self.configure(fg_color=COR_FUNDO)
        self.on_save = on_save

        ctk.CTkLabel(
            self,
            text="Chave de API da Anthropic",
            font=FONT_PASSO, text_color=COR_TEXTO, anchor="w",
        ).pack(padx=32, pady=(24, 4), anchor="w")

        ctk.CTkLabel(
            self,
            text="Acesse console.anthropic.com → API Keys e cole a chave abaixo.",
            font=FONT_AJUDA, text_color=COR_TEXTO_FRACO,
            justify="left", anchor="w",
        ).pack(padx=32, pady=(0, 14), anchor="w")

        self.entry = ctk.CTkEntry(
            self, show="•", height=ALTURA_ENTRY, font=FONT_ENTRY,
            border_color=COR_BORDA, text_color=COR_TEXTO,
            placeholder_text="sk-ant-…",
        )
        self.entry.pack(fill="x", padx=32, pady=(0, 16))
        if current_key:
            self.entry.insert(0, current_key)

        ctk.CTkButton(
            self, text="SALVAR E FECHAR", height=ALTURA_BTN,
            font=FONT_BTN, corner_radius=10,
            fg_color=COR_ACAO, hover_color=COR_ACAO_HOVER,
            command=self._save,
        ).pack(padx=32, pady=(0, 24), fill="x")

    def _save(self):
        k = self.entry.get().strip()
        if not k:
            messagebox.showwarning("Chave em branco",
                                   "Digite a chave API antes de salvar.", parent=self)
            return
        self.on_save(k)
        self.destroy()

# ───────────────────────────────────────────────────────────────────────────────
# APLICAÇÃO PRINCIPAL
# ───────────────────────────────────────────────────────────────────────────────

PAD = {"padx": 30, "pady": 7}

class App(ctk.CTk):

    def __init__(self):
        super().__init__()
        self.config_data  = load_config()
        self.api_key      = self.config_data.get("api_key", "")

        self.processo_folder = ctk.StringVar()
        self.pre_laudo_path  = ctk.StringVar()
        self.campo_path      = ctk.StringVar()
        self.photos_sub      = ctk.StringVar()
        self.photos_list: list     = []
        self.avaliacoes_paths: list = []

        # Identificação do processo
        self.reclamante     = ctk.StringVar()
        self.reclamada      = ctk.StringVar()
        self.tipo_insalubr  = ctk.BooleanVar(value=True)
        self.tipo_periculos = ctk.BooleanVar(value=False)
        self.funcao         = ctk.StringVar()

        self.imp_doc_path   = ctk.StringVar()
        self.meu_laudo_path = ctk.StringVar()

        self._pausar_trace = False

        # Auto-preenchimento ao selecionar documentos manualmente
        self.pre_laudo_path.trace_add('write', self._on_doc_changed)
        self.campo_path.trace_add('write', self._on_doc_changed)

        ctk.set_appearance_mode("light")
        ctk.set_default_color_theme("blue")
        self._build_ui()
        self.after(50, lambda: self.state("zoomed"))

        if not self.api_key:
            self.after(600, self._open_settings)

    # ── Construção da interface ─────────────────────────────────────────────

    def _build_ui(self):
        self.title("SISTEMA ARI — Laudos Periciais")
        self.resizable(True, True)
        self.configure(fg_color=COR_FUNDO)

        # ── Cabeçalho ──────────────────────────────────────────────────────
        hdr = ctk.CTkFrame(self, fg_color=COR_FUNDO_HDR, corner_radius=0, height=96)
        hdr.pack(fill="x")
        hdr.pack_propagate(False)

        hdr_left = ctk.CTkFrame(hdr, fg_color="transparent")
        hdr_left.pack(side="left", padx=32, fill="y")

        ctk.CTkLabel(
            hdr_left, text="SISTEMA ARI",
            font=FONT_TITULO, text_color="white", anchor="w",
        ).pack(anchor="w", pady=(18, 0))
        ctk.CTkLabel(
            hdr_left, text="Laudos Periciais  ·  Ari V. Copesco Jr.  ·  CREA 060097553-3",
            font=FONT_SUBTITULO, text_color="#b6c4d8", anchor="w",
        ).pack(anchor="w")

        ctk.CTkButton(
            hdr, text="⚙   Configurar API", width=200, height=48,
            font=FONT_BTN, fg_color=COR_SECUNDARIO, hover_color=COR_SEC_HOVER,
            corner_radius=8, command=self._open_settings,
        ).pack(side="right", padx=32)

        # ── PASSO 1 — Pasta do Processo ────────────────────────────────────
        p1 = self._cartao(self, padx=24, pady=(20, 10))
        self._titulo_passo(p1, "1", "Pasta do Processo",
                            "Selecione a pasta onde estão o pré-laudo, as fotos e os documentos.")

        row = ctk.CTkFrame(p1, fg_color="transparent")
        row.pack(fill="x", padx=24, pady=(4, 8))
        ctk.CTkEntry(
            row, textvariable=self.processo_folder,
            height=ALTURA_ENTRY, font=FONT_ENTRY,
            border_color=COR_BORDA, text_color=COR_TEXTO,
        ).pack(side="left", fill="x", expand=True, padx=(0, 12))
        ctk.CTkButton(
            row, text="ESCOLHER PASTA", width=220, height=ALTURA_ENTRY,
            font=FONT_BTN, fg_color=COR_ACAO, hover_color=COR_ACAO_HOVER,
            corner_radius=8, command=self._pick_processo,
        ).pack(side="right")

        self.detect_lbl = ctk.CTkLabel(
            p1, text="Nenhuma pasta selecionada ainda.",
            font=FONT_AJUDA, text_color=COR_TEXTO_FRACO, anchor="w",
        )
        self.detect_lbl.pack(fill="x", padx=24, pady=(0, 18))

        # ── Abas ────────────────────────────────────────────────────────────
        tabs = ctk.CTkTabview(
            self, fg_color=COR_FUNDO,
            segmented_button_fg_color=COR_DESTAQUE,
            segmented_button_selected_color=COR_ACAO,
            segmented_button_selected_hover_color=COR_ACAO_HOVER,
            segmented_button_unselected_color=COR_DESTAQUE,
            segmented_button_unselected_hover_color="#c5cdd9",
            text_color="white",
            text_color_disabled=COR_TEXTO,
        )
        # Aumenta a altura dos botões do segmento
        try:
            tabs._segmented_button.configure(height=52, font=FONT_BTN)
        except Exception:
            pass

        tabs.pack(fill="both", expand=True, padx=20, pady=(8, 20))
        tabs.add("   Gerar Laudo Pericial   ")
        tabs.add("   Responder Impugnação / Quesitos   ")

        self._build_laudo_tab(tabs.tab("   Gerar Laudo Pericial   "))
        self._build_resposta_tab(tabs.tab("   Responder Impugnação / Quesitos   "))

    # ── Aba: Gerar Laudo ─────────────────────────────────────────────────────

    def _build_laudo_tab(self, parent):
        # Wrapper com grid: linha 0 = scroll (expande), linha 1 = botão fixo
        wrapper = ctk.CTkFrame(parent, fg_color=COR_FUNDO, corner_radius=0)
        wrapper.pack(fill="both", expand=True)
        wrapper.grid_rowconfigure(0, weight=1)
        wrapper.grid_rowconfigure(1, weight=0)
        wrapper.grid_columnconfigure(0, weight=1)

        # ── Área rolável com os passos ──────────────────────────────────────
        main = ctk.CTkScrollableFrame(wrapper, fg_color=COR_FUNDO)
        main.grid(row=0, column=0, sticky="nsew")

        # Passo 2 — Identificação do Processo
        c2 = self._cartao(main)
        self._titulo_passo(c2, "2", "Identificação do Processo",
                            "Preenchido automaticamente ao escolher a pasta.")

        row_rec = ctk.CTkFrame(c2, fg_color="transparent")
        row_rec.pack(fill="x", padx=24, pady=(4, 8))
        for label, var, w in [
            ("Reclamante", self.reclamante, 1),
            ("Reclamada",  self.reclamada,  1),
        ]:
            col = ctk.CTkFrame(row_rec, fg_color="transparent")
            col.pack(side="left", fill="x", expand=True,
                     padx=(0, 16) if label == "Reclamante" else (0, 0))
            ctk.CTkLabel(col, text=label, font=FONT_LABEL,
                         text_color=COR_TEXTO, anchor="w").pack(anchor="w")
            ctk.CTkEntry(col, textvariable=var, height=ALTURA_ENTRY,
                         font=FONT_ENTRY, border_color=COR_BORDA,
                         text_color=COR_TEXTO).pack(fill="x", pady=(2, 0))

        row_fc = ctk.CTkFrame(c2, fg_color="transparent")
        row_fc.pack(fill="x", padx=24, pady=(0, 8))

        col_fn = ctk.CTkFrame(row_fc, fg_color="transparent")
        col_fn.pack(side="left", fill="x", expand=True, padx=(0, 16))
        ctk.CTkLabel(col_fn, text="Função / Cargo", font=FONT_LABEL,
                     text_color=COR_TEXTO, anchor="w").pack(anchor="w")
        ctk.CTkEntry(col_fn, textvariable=self.funcao, height=ALTURA_ENTRY,
                     font=FONT_ENTRY, border_color=COR_BORDA,
                     text_color=COR_TEXTO).pack(fill="x", pady=(2, 0))

        col_tp = ctk.CTkFrame(row_fc, fg_color="transparent")
        col_tp.pack(side="left", fill="none")
        ctk.CTkLabel(col_tp, text="Tipo de Perícia", font=FONT_LABEL,
                     text_color=COR_TEXTO, anchor="w").pack(anchor="w")
        row_chk = ctk.CTkFrame(col_tp, fg_color="transparent")
        row_chk.pack(anchor="w", pady=(6, 0))
        ctk.CTkCheckBox(row_chk, text=" Insalubridade",
                        variable=self.tipo_insalubr,
                        font=FONT_LABEL, text_color=COR_TEXTO,
                        fg_color=COR_ACAO, hover_color=COR_ACAO_HOVER,
                        checkbox_width=24, checkbox_height=24,
                        ).pack(side="left", padx=(0, 20))
        ctk.CTkCheckBox(row_chk, text=" Periculosidade",
                        variable=self.tipo_periculos,
                        font=FONT_LABEL, text_color=COR_TEXTO,
                        fg_color=COR_ACAO, hover_color=COR_ACAO_HOVER,
                        checkbox_width=24, checkbox_height=24,
                        ).pack(side="left")

        self.ident_status = ctk.CTkLabel(
            c2, text="Aguardando seleção da pasta…",
            font=FONT_AJUDA, text_color=COR_TEXTO_FRACO, anchor="w",
        )
        self.ident_status.pack(fill="x", padx=24, pady=(6, 14))

        # Passo 3 — Pré-laudo
        c3 = self._cartao(main)
        self._titulo_passo(c3, "3", "Pré-Laudo",
                            "Documento preparado pelas secretárias.")
        self._file_row(c3, self.pre_laudo_path,
                       [("Word / PDF", "*.docx *.pdf")],
                       lambda: self.processo_folder.get() or None)

        # Passo 4 — Campo
        c4 = self._cartao(main)
        self._titulo_passo(c4, "4", "Anotações de Campo",
                            "Dados coletados durante a diligência.")
        self._file_row(c4, self.campo_path,
                       [("Word / PDF", "*.docx *.pdf")],
                       lambda: self.processo_folder.get() or None)

        # Passo 5 — Fotos
        c5 = self._cartao(main)
        self._titulo_passo(c5, "5", "Fotos da Diligência",
                            "Subpasta onde estão as imagens.")
        row5 = ctk.CTkFrame(c5, fg_color="transparent")
        row5.pack(fill="x", padx=24, pady=(4, 8))
        ctk.CTkEntry(
            row5, textvariable=self.photos_sub,
            height=ALTURA_ENTRY, font=FONT_ENTRY,
            border_color=COR_BORDA, text_color=COR_TEXTO,
        ).pack(side="left", fill="x", expand=True, padx=(0, 12))
        ctk.CTkButton(
            row5, text="ESCOLHER SUBPASTA", width=220, height=ALTURA_ENTRY,
            font=FONT_BTN, fg_color=COR_SECUNDARIO, hover_color=COR_SEC_HOVER,
            corner_radius=8, command=self._pick_photos_sub,
        ).pack(side="right")

        self.photos_lbl = ctk.CTkLabel(
            c5, text="Nenhuma subpasta selecionada.",
            font=FONT_AJUDA, text_color=COR_TEXTO_FRACO, anchor="w",
        )
        self.photos_lbl.pack(fill="x", padx=24, pady=(0, 14))

        # Passo 6 — Avaliações Técnicas
        c6 = self._cartao(main)
        self._titulo_passo(c6, "6", "Avaliações Técnicas",
                            "Detectadas automaticamente na pasta do processo. "
                            "Use os botões abaixo para incluir/remover manualmente.")

        row6 = ctk.CTkFrame(c6, fg_color="transparent")
        row6.pack(fill="x", padx=24, pady=(4, 6))
        ctk.CTkButton(
            row6, text="ADICIONAR ARQUIVO(S)", width=240, height=ALTURA_BTN - 8,
            font=FONT_BTN, fg_color=COR_SECUNDARIO, hover_color=COR_SEC_HOVER,
            corner_radius=8, command=self._pick_avaliacoes,
        ).pack(side="left", padx=(0, 12))
        ctk.CTkButton(
            row6, text="LIMPAR", width=120, height=ALTURA_BTN - 8,
            font=FONT_BTN, fg_color="#7a8190", hover_color="#5a6270",
            corner_radius=8, command=self._clear_avaliacoes,
        ).pack(side="left")

        self.aval_lbl = ctk.CTkLabel(
            c6, text="Nenhuma avaliação técnica adicionada.",
            font=FONT_AJUDA, text_color=COR_TEXTO_FRACO, anchor="w",
        )
        self.aval_lbl.pack(fill="x", padx=24, pady=(0, 14))

        # Passo 7 — Observações
        c7 = self._cartao(main)
        self._titulo_passo(c7, "7", "Observações",
                            "Opcional — instruções, medições ou pontos a destacar.")

        self.obs_laudo = ctk.CTkTextbox(
            c7, height=160, font=FONT_ENTRY,
            fg_color=COR_OBS_BG, border_width=2, border_color=COR_OBS_BORDA,
            text_color=COR_TEXTO,
        )
        self.obs_laudo.pack(fill="x", padx=24, pady=(4, 16))

        # ── Botão fixo (linha 1 do grid — sempre visível) ──────────────────
        bottom = ctk.CTkFrame(wrapper, fg_color=COR_FUNDO, corner_radius=0)
        bottom.grid(row=1, column=0, sticky="ew")

        self.btn_laudo = ctk.CTkButton(
            bottom, text="▶   GERAR LAUDO PERICIAL", height=ALTURA_BTN_GR,
            font=FONT_BTN_GR, corner_radius=12,
            fg_color=COR_ACAO, hover_color=COR_ACAO_HOVER,
            command=self._start_laudo,
        )
        self.btn_laudo.pack(fill="x", padx=24, pady=(10, 4))

        self.prog_laudo = ctk.CTkProgressBar(
            bottom, mode="indeterminate", height=14,
            fg_color=COR_DESTAQUE, progress_color=COR_ACAO,
        )
        self.prog_laudo.pack(fill="x", padx=24)
        self.prog_laudo.set(0)

        self.lbl_laudo = ctk.CTkLabel(
            bottom, text="Pronto para gerar.",
            font=FONT_STATUS, text_color=COR_TEXTO_FRACO,
        )
        self.lbl_laudo.pack(pady=(4, 10))

    # ── Aba: Responder Impugnação ────────────────────────────────────────────

    def _build_resposta_tab(self, parent):
        main = ctk.CTkScrollableFrame(parent, fg_color=COR_FUNDO)
        main.pack(fill="both", expand=True)

        # Passo 3 — Documento recebido
        c3 = self._cartao(main)
        self._titulo_passo(c3, "3", "Documento Recebido",
                            "Impugnação ou quesitos complementares enviados pela parte.")
        self._file_row(c3, self.imp_doc_path,
                       [("Word / PDF", "*.docx *.pdf")],
                       lambda: self.processo_folder.get() or None)

        # Passo 4 — Meu laudo
        c4 = self._cartao(main)
        self._titulo_passo(c4, "4", "Meu Laudo",
                            "Laudo original (referência para defender as conclusões).")
        row4 = ctk.CTkFrame(c4, fg_color="transparent")
        row4.pack(fill="x", padx=24, pady=(4, 6))
        ctk.CTkEntry(
            row4, textvariable=self.meu_laudo_path,
            height=ALTURA_ENTRY, font=FONT_ENTRY,
            border_color=COR_BORDA, text_color=COR_TEXTO,
        ).pack(side="left", fill="x", expand=True, padx=(0, 12))
        ctk.CTkButton(
            row4, text="ESCOLHER ARQUIVO", width=220, height=ALTURA_ENTRY,
            font=FONT_BTN, fg_color=COR_SECUNDARIO, hover_color=COR_SEC_HOVER,
            corner_radius=8,
            command=lambda: self._pick_file(
                self.meu_laudo_path,
                [("Word / PDF", "*.docx *.pdf")],
                self.processo_folder.get() or None),
        ).pack(side="right")

        self.laudo_ref_lbl = ctk.CTkLabel(
            c4,
            text="Auto-detectado ao selecionar a pasta do processo.",
            font=FONT_AJUDA, text_color=COR_TEXTO_FRACO, anchor="w",
        )
        self.laudo_ref_lbl.pack(fill="x", padx=24, pady=(0, 14))

        # Passo 5 — Observações
        c5 = self._cartao(main)
        self._titulo_passo(c5, "5", "Observações",
                            "Opcional — pontos específicos a destacar na resposta.")
        self.obs_resp = ctk.CTkTextbox(
            c5, height=160, font=FONT_ENTRY,
            fg_color=COR_OBS_BG, border_width=2, border_color=COR_OBS_BORDA,
            text_color=COR_TEXTO,
        )
        self.obs_resp.pack(fill="x", padx=24, pady=(4, 16))

        # Botão grande
        self.btn_resp = ctk.CTkButton(
            main, text="▶   GERAR RESPOSTA / ESCLARECIMENTOS",
            height=ALTURA_BTN_GR, font=FONT_BTN_GR, corner_radius=12,
            fg_color=COR_ACAO, hover_color=COR_ACAO_HOVER,
            command=self._start_resposta,
        )
        self.btn_resp.pack(padx=24, pady=(8, 6), fill="x")

        self.prog_resp = ctk.CTkProgressBar(
            main, mode="indeterminate", height=16,
            fg_color=COR_DESTAQUE, progress_color=COR_ACAO,
        )
        self.prog_resp.pack(padx=24, fill="x")
        self.prog_resp.set(0)

        self.lbl_resp = ctk.CTkLabel(
            main, text="Pronto para gerar.", font=FONT_STATUS, text_color=COR_TEXTO_FRACO,
        )
        self.lbl_resp.pack(pady=(10, 32))

    # ── Helpers de UI ────────────────────────────────────────────────────────

    def _cartao(self, parent, padx=24, pady=(0, 12)):
        """Cartão branco com borda suave que agrupa um passo do fluxo."""
        card = ctk.CTkFrame(
            parent, fg_color=COR_FUNDO_CARD, corner_radius=12,
            border_width=1, border_color=COR_BORDA,
        )
        card.pack(fill="x", padx=padx, pady=pady)
        return card

    def _titulo_passo(self, parent, num: str, titulo: str, ajuda: str = ""):
        """Cabeçalho de cartão: bolinha azul com número + título + descrição."""
        wrap = ctk.CTkFrame(parent, fg_color="transparent")
        wrap.pack(fill="x", padx=24, pady=(18, 6))

        bola = ctk.CTkFrame(
            wrap, fg_color=COR_ACAO, corner_radius=22, width=44, height=44,
        )
        bola.pack(side="left", padx=(0, 16))
        bola.pack_propagate(False)
        ctk.CTkLabel(
            bola, text=num, font=FONT_PASSO_NUM, text_color="white",
        ).pack(expand=True)

        txt = ctk.CTkFrame(wrap, fg_color="transparent")
        txt.pack(side="left", fill="x", expand=True)
        ctk.CTkLabel(
            txt, text=titulo, font=FONT_PASSO, text_color=COR_TEXTO, anchor="w",
        ).pack(anchor="w")
        if ajuda:
            ctk.CTkLabel(
                txt, text=ajuda, font=FONT_AJUDA, text_color=COR_TEXTO_FRACO,
                anchor="w", justify="left",
            ).pack(anchor="w", pady=(2, 0))

    def _file_row(self, parent, var, filetypes, get_dir=None):
        row = ctk.CTkFrame(parent, fg_color="transparent")
        row.pack(fill="x", padx=24, pady=(4, 14))
        ctk.CTkEntry(
            row, textvariable=var,
            height=ALTURA_ENTRY, font=FONT_ENTRY,
            border_color=COR_BORDA, text_color=COR_TEXTO,
        ).pack(side="left", fill="x", expand=True, padx=(0, 12))
        ctk.CTkButton(
            row, text="ESCOLHER ARQUIVO", width=220, height=ALTURA_ENTRY,
            font=FONT_BTN, fg_color=COR_SECUNDARIO, hover_color=COR_SEC_HOVER,
            corner_radius=8,
            command=lambda: self._pick_file(var, filetypes,
                                            get_dir() if get_dir else None),
        ).pack(side="right")

    def _pick_file(self, var, filetypes, initial_dir=None):
        kw = {"initialdir": initial_dir} if initial_dir and Path(initial_dir).is_dir() else {}
        p = filedialog.askopenfilename(filetypes=filetypes, **kw)
        if p:
            var.set(p)

    def _pick_processo(self):
        folder = filedialog.askdirectory()
        if not folder:
            return
        self.processo_folder.set(folder)
        det = auto_detect(folder)

        # ── Preenche caminhos detectados ───────────────────────────────────
        # Desativa trace temporariamente para não disparar _on_doc_changed
        # antes de termos todos os dados da pasta.
        self._pausar_trace = True
        if det["pre_laudo"]:
            self.pre_laudo_path.set(det["pre_laudo"])
        if det["campo"]:
            self.campo_path.set(det["campo"])
        self._pausar_trace = False

        if det["photos_sub"]:
            self.photos_sub.set(det["photos_sub"])
            self.photos_list = get_photos(det["photos_sub"])
            n = min(det["photo_count"], MAX_PHOTOS)
            self.photos_lbl.configure(
                text=f"✓  {n} foto(s) na subpasta '{Path(det['photos_sub']).name}'.",
                text_color=COR_OK,
            )
        if det["laudo"]:
            self.meu_laudo_path.set(det["laudo"])
            self.laudo_ref_lbl.configure(
                text=f"✓  Detectado: {Path(det['laudo']).name}",
                text_color=COR_OK,
            )
        if det["avaliacoes"]:
            self.avaliacoes_paths = list(det["avaliacoes"])
            self._update_aval_label()

        partes = [
            "Pré-laudo: " + ("✓" if det["pre_laudo"] else "—"),
            "Campo: "     + ("✓" if det["campo"]     else "—"),
            f"Fotos: {min(det['photo_count'], MAX_PHOTOS)}" if det["photo_count"] else "Fotos: —",
            "Laudo anterior: " + ("✓" if det["laudo"] else "—"),
            f"Avaliações: {len(det['avaliacoes'])}" if det["avaliacoes"] else "Avaliações: —",
        ]
        self.detect_lbl.configure(
            text="   ".join(partes),
            text_color=COR_TEXTO,
        )

        # ── Lê dados do processo (pasta + documentos) ──────────────────────
        self._auto_fill_identificacao(folder, det)

    def _pick_photos_sub(self):
        folder = filedialog.askdirectory(
            initialdir=self.processo_folder.get() or None
        )
        if not folder:
            return
        self.photos_sub.set(folder)
        self.photos_list = get_photos(folder)
        n = len(self.photos_list)
        self.photos_lbl.configure(
            text=(f"✓  {n} foto(s) encontrada(s)."
                  + (f"  (limite {MAX_PHOTOS} ativas)" if n >= MAX_PHOTOS else "")
                  if n else "⚠  Nenhuma foto encontrada nesta subpasta."),
            text_color=COR_OK if n else COR_AVISO,
        )

    def _pick_avaliacoes(self):
        kw = {}
        d = self.processo_folder.get()
        if d and Path(d).is_dir():
            kw["initialdir"] = d
        files = filedialog.askopenfilenames(
            filetypes=[("Word / PDF", "*.docx *.pdf")], **kw
        )
        for f in files:
            p = Path(f)
            if p not in self.avaliacoes_paths:
                self.avaliacoes_paths.append(p)
        self._update_aval_label()

    def _clear_avaliacoes(self):
        self.avaliacoes_paths = []
        self._update_aval_label()

    def _update_aval_label(self):
        n = len(self.avaliacoes_paths)
        if n == 0:
            self.aval_lbl.configure(
                text="Nenhuma avaliação técnica adicionada.",
                text_color=COR_TEXTO_FRACO,
            )
        else:
            names = ", ".join(p.name for p in self.avaliacoes_paths[:3])
            if n > 3:
                names += f" (+{n - 3})"
            self.aval_lbl.configure(
                text=f"✓  {n} arquivo(s): {names}", text_color=COR_OK,
            )

    def _set_sl(self, msg, color=None):
        self.lbl_laudo.configure(text=msg, text_color=color or COR_TEXTO_FRACO)
        self.update_idletasks()

    def _set_sr(self, msg, color=None):
        self.lbl_resp.configure(text=msg, text_color=color or COR_TEXTO_FRACO)
        self.update_idletasks()

    def _open_settings(self):
        SettingsDialog(self, self.api_key, self._save_key)

    def _save_key(self, key: str):
        self.api_key = key
        self.config_data["api_key"] = key
        save_config(self.config_data)

    def _on_doc_changed(self, *_):
        """Chamado pelo trace quando pré-laudo ou campo é trocado manualmente."""
        if getattr(self, '_pausar_trace', False):
            return
        folder = self.processo_folder.get()
        det = {
            'pre_laudo': self.pre_laudo_path.get() or None,
            'campo':     self.campo_path.get() or None,
        }
        self._auto_fill_identificacao(folder or '', det)

    def _auto_fill_identificacao(self, folder: str, det: dict):
        """Extrai e preenche Reclamante, Reclamada, Função e tipo."""
        try:
            data = parse_processo_da_pasta(folder, det)
        except Exception:
            data = {}

        filled = []
        if data.get('reclamante'):
            self.reclamante.set(data['reclamante'])
            filled.append('Reclamante')
        if data.get('reclamada'):
            self.reclamada.set(data['reclamada'])
            filled.append('Reclamada')
        if data.get('funcao'):
            self.funcao.set(data['funcao'])
            filled.append('Função')

        insalubr  = bool(data.get('insalubr',  False))
        periculos = bool(data.get('periculos', False))
        if not insalubr and not periculos:
            insalubr = True   # default: insalubridade
        self.tipo_insalubr.set(insalubr)
        self.tipo_periculos.set(periculos)
        if insalubr or periculos:
            tipo = ' + '.join(filter(None, [
                'Insalubr.' if insalubr else '',
                'Periculos.' if periculos else '',
            ]))
            filled.append(f'Tipo: {tipo}')

        origem = ''
        if det.get('pre_laudo'):
            origem = Path(det['pre_laudo']).name
        elif folder:
            origem = Path(folder).name

        if filled:
            self.ident_status.configure(
                text=f"✓  Preenchido automaticamente: {', '.join(filled)}"
                     + (f"  —  {origem}" if origem else ""),
                text_color=COR_OK,
            )
        else:
            self.ident_status.configure(
                text="⚠  Não foi possível extrair dados. Preencha manualmente.",
                text_color=COR_AVISO,
            )

    # ── Gerar Laudo ──────────────────────────────────────────────────────────

    def _start_laudo(self):
        erros = []
        if not self.api_key:
            erros.append("- Chave API não configurada (clique em ⚙ Configurar API)")
        if not self.processo_folder.get():
            erros.append("- Pasta do processo não selecionada (Passo 1)")
        if not self.pre_laudo_path.get():
            erros.append("- Pré-laudo não selecionado (Passo 2)")
        if not self.campo_path.get():
            erros.append("- Anotações de campo não selecionadas (Passo 3)")
        if erros:
            messagebox.showerror("Campos obrigatórios", "\n".join(erros))
            return
        self.btn_laudo.configure(state="disabled")
        self.prog_laudo.start()
        self._set_sl("Iniciando...", COR_ACAO)
        threading.Thread(target=self._laudo_thread, daemon=True).start()

    def _laudo_thread(self):
        try:
            self._set_sl("Lendo pré-laudo...")
            pre  = extract_text(self.pre_laudo_path.get())
            self._set_sl("Lendo anotações de campo...")
            camp = extract_text(self.campo_path.get())

            avaliacoes_text = ""
            if self.avaliacoes_paths:
                self._set_sl("Lendo avaliações técnicas...")
                parts = []
                for p in self.avaliacoes_paths:
                    txt = extract_text(str(p))
                    parts.append(f"=== {p.name} ===\n{txt}")
                avaliacoes_text = "\n\n".join(parts)

            self._set_sl("Carregando perfil de escrita e laudos de referência...")
            agente, perfil = load_profile_files()
            ref = load_reference_laudo()
            obs = self.obs_laudo.get("1.0", "end").strip()

            fname = build_laudo_filename(
                self.reclamante.get(), self.reclamada.get(),
                self.tipo_insalubr.get(), self.tipo_periculos.get(),
                self.funcao.get()
            )
            out = output_path(self.processo_folder.get(), fname)

            txt = gerar_laudo(
                api_key=self.api_key, pre_laudo=pre, campo=camp,
                photos=self.photos_list, agente_md=agente, perfil_md=perfil,
                ref_laudo=ref, obs=obs, avaliacoes=avaliacoes_text,
                progress_cb=self._set_sl,
            )

            self._set_sl("Salvando .docx...")
            save_docx(txt, out)

            self.meu_laudo_path.set(out)
            self.laudo_ref_lbl.configure(
                text=f"  {Path(out).name}  (gerado agora)", text_color=COR_OK
            )

            self.prog_laudo.stop()
            self.prog_laudo.set(1)
            self._set_sl(f"  Salvo: {Path(out).name}", COR_OK)
            self.btn_laudo.configure(state="normal")

            pasta_final = Path(out).parent
            abrir = messagebox.askyesno(
                "Laudo gerado!",
                f"Laudo salvo com sucesso!\n\n"
                f"Arquivo: {Path(out).name}\n"
                f"Pasta: {pasta_final}\n\n"
                f"Deseja abrir a pasta agora?"
            )
            if abrir:
                abrir_pasta_no_explorer(pasta_final)
        except anthropic.AuthenticationError:
            self.prog_laudo.stop()
            self.btn_laudo.configure(state="normal")
            self._set_sl("  Chave API inválida.", COR_ERRO)
            messagebox.showerror("Erro de autenticação",
                                 "Chave API inválida.\nClique em ⚙ Configurar API.")
        except Exception as e:
            self.prog_laudo.stop()
            self.btn_laudo.configure(state="normal")
            self._set_sl(f"  Erro: {e}", COR_ERRO)
            messagebox.showerror("Erro", f"Ocorreu um erro:\n\n{e}")

    # ── Gerar Resposta ───────────────────────────────────────────────────────

    def _start_resposta(self):
        erros = []
        if not self.api_key:
            erros.append("- Chave API não configurada (clique em ⚙ Configurar API)")
        if not self.processo_folder.get():
            erros.append("- Pasta do processo não selecionada (Passo 1)")
        if not self.imp_doc_path.get():
            erros.append("- Documento recebido não selecionado (Passo 2)")
        if erros:
            messagebox.showerror("Campos obrigatórios", "\n".join(erros))
            return
        self.btn_resp.configure(state="disabled")
        self.prog_resp.start()
        self._set_sr("Iniciando...", COR_ACAO)
        threading.Thread(target=self._resposta_thread, daemon=True).start()

    def _resposta_thread(self):
        try:
            self._set_sr("Lendo documento recebido...")
            doc  = extract_text(self.imp_doc_path.get())
            laudo = ""
            if self.meu_laudo_path.get():
                self._set_sr("Lendo meu laudo original...")
                laudo = extract_text(self.meu_laudo_path.get())
            self._set_sr("Carregando perfil de escrita...")
            agente, perfil = load_profile_files()
            obs = self.obs_resp.get("1.0", "end").strip()

            today = date.today().strftime("%Y%m%d")
            out = output_path(self.processo_folder.get(), f"ESCLARECIMENTOS_{today}")

            txt = gerar_resposta(
                api_key=self.api_key, doc_recebido=doc, meu_laudo=laudo,
                agente_md=agente, perfil_md=perfil,
                obs=obs, progress_cb=self._set_sr,
            )

            self._set_sr("Salvando .docx...")
            save_docx(txt, out)

            self.prog_resp.stop()
            self.prog_resp.set(1)
            self._set_sr(f"  Salvo: {Path(out).name}", COR_OK)
            self.btn_resp.configure(state="normal")
            messagebox.showinfo(
                "Concluído",
                f"Resposta gerada com sucesso!\n\n"
                f"Arquivo: {Path(out).name}\nPasta: {self.processo_folder.get()}"
            )
        except anthropic.AuthenticationError:
            self.prog_resp.stop()
            self.btn_resp.configure(state="normal")
            self._set_sr("  Chave API inválida.", COR_ERRO)
            messagebox.showerror("Erro de autenticação",
                                 "Chave API inválida.\nClique em ⚙ Configurar API.")
        except Exception as e:
            self.prog_resp.stop()
            self.btn_resp.configure(state="normal")
            self._set_sr(f"  Erro: {e}", COR_ERRO)
            messagebox.showerror("Erro", f"Ocorreu um erro:\n\n{e}")


# ───────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = App()
    app.mainloop()

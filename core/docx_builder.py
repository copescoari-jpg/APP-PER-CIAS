"""
Geração do .docx no padrão do perito Ari Vladimir Copesco Junior.
Extraído e centralizado aqui para uso por main.py.
"""

import io
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Cabeçalhos de seção reconhecidos ──────────────────────────────────────────

SECTION_HEADERS = [
    "OBJETIVO", "PROCESSO", "METODOLOGIA", "DILIGÊNCIAS",
    "ATIVIDADES DA RECLAMANTE", "ATIVIDADES DO RECLAMANTE",
    "DESCRIÇÃO DO LOCAL", "EQUIPAMENTOS DE PROTEÇÃO INDIVIDUAL",
    "RESULTADOS DAS AVALIAÇÕES", "RESPOSTAS AOS QUESITOS",
    "HONORÁRIOS PERICIAIS", "CONCLUSÃO", "ENCERRAMENTO",
    "Quesitos da Reclamante", "Quesitos da Reclamada",
    "Quesitos do Juízo", "Quesitos Complementares",
    "ESCLARECIMENTOS", "CONSIDERAÇÕES",
    "IDENTIFICAÇÃO DO PROCESSO", "LOCAL DE TRABALHO", "VERSÃO DO RECLAMANTE",
    "VERSÃO DA RECLAMADA", "MEDIÇÕES E OBSERVAÇÕES", "EPI FORNECIDO",
    "FUNDAMENTAÇÃO", "RESPOSTAS AOS QUESITOS",
    "ANEXOS", "ANEXOS — AVALIAÇÕES TÉCNICAS",
]
_NUMBERED = {"6", "7", "8", "9", "10", "11", "12"}

_RE_RESPOSTA  = re.compile(r'^(Resposta\s*:)(.*)', re.IGNORECASE | re.DOTALL)
_RE_QUESITO_N = re.compile(r'^(\d{1,2}\s*[-\.\)]\s+|[a-zA-Z][\.\)]\s+)(.+)', re.DOTALL)


def _is_header(line: str) -> bool:
    u = line.strip().upper()
    for n in _NUMBERED:
        if u.startswith(f"{n} -") or u.startswith(f"{n}-") or u.startswith(f"{n}."):
            return True
    return any(u == h.upper() or u.startswith(h.upper()) for h in SECTION_HEADERS)


def _footer_add_page_number(paragraph):
    """Insere campo PAGE no parágrafo: — X —"""
    def _field_run(fld_type=None, instr=None):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:ascii'), 'Arial')
        rf.set(qn('w:hAnsi'), 'Arial')
        rPr.append(rf)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '18')
        rPr.append(sz)
        rPr.append(OxmlElement('w:i'))
        r.append(rPr)
        if fld_type:
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), fld_type)
            r.append(fc)
        if instr:
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = instr
            r.append(it)
        return r

    def _txt(txt):
        run = paragraph.add_run(txt)
        run.font.name   = "Arial"
        run.font.size   = Pt(9)
        run.font.italic = True

    _txt(" — ")
    paragraph._p.append(_field_run(fld_type='begin'))
    paragraph._p.append(_field_run(instr=' PAGE '))
    paragraph._p.append(_field_run(fld_type='end'))
    _txt(" —")


def _add_header_footer(sec):
    """Cabeçalho LAUDO TÉCNICO + rodapé com endereço e número de página."""
    hdr = sec.header
    p0 = hdr.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p1 = hdr.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(4)
    run1 = p1.add_run("LAUDO TÉCNICO")
    run1.font.name = "Arial"
    run1.font.size = Pt(18)
    run1.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    p2 = hdr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p2._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '32')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    ftr = sec.footer
    pPr0f = ftr.paragraphs[0]._p.get_or_add_pPr()
    pBdr0f = OxmlElement('w:pBdr')
    top_bdr = OxmlElement('w:top')
    top_bdr.set(qn('w:val'), 'single')
    top_bdr.set(qn('w:sz'), '8')
    top_bdr.set(qn('w:space'), '1')
    top_bdr.set(qn('w:color'), '808080')
    pBdr0f.append(top_bdr)
    pPr0f.append(pBdr0f)

    footer_lines = [
        ("Ari Vladimir Copesco Júnior",                             10),
        ("Engenheiro Segurança do Trabalho CREA 060097553-3",       9),
        ("Rua Marechal Rondon, n° 224  ·  Fone (016) 3235-6763",   9),
        ("Ribeirão Preto/SP  ·  CEP 14.025-430",                   9),
    ]
    for i, (line, size) in enumerate(footer_lines):
        p = ftr.paragraphs[0] if i == 0 else ftr.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name   = "Arial"
        run.font.size   = Pt(size)
        run.font.italic = True

    p_num = ftr.add_paragraph()
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_num.paragraph_format.space_before = Pt(3)
    p_num.paragraph_format.space_after  = Pt(0)
    _footer_add_page_number(p_num)


def salvar_docx(texto: str, caminho: str, avaliacoes_paths: list | None = None):
    """
    Gera e salva o laudo como .docx no caminho indicado.

    Args:
        texto: Texto plano do laudo (separado por \\n).
        caminho: Caminho completo de saída (.docx).
        avaliacoes_paths: Lista de Paths de PDFs de avaliação para anexar.
    """
    from core.pdf_reader import pdf_para_imagens

    doc = Document()
    for sec in doc.sections:
        sec.top_margin      = Emu(540385)
        sec.bottom_margin   = Emu(360045)
        sec.left_margin     = Emu(1080135)
        sec.right_margin    = Emu(1080135)
        sec.header_distance = Emu(180340)
        sec.footer_distance = Emu(360045)
        _add_header_footer(sec)

    sty = doc.styles["Normal"]
    sty.font.name = "Arial"
    sty.font.size = Pt(12)
    sty.paragraph_format.line_spacing = 1.15

    for raw in texto.split("\n"):
        s = raw.strip()
        if not s:
            doc.add_paragraph("")
            continue

        p = doc.add_paragraph()

        def _run(txt, bold=False, italic=False):
            r = p.add_run(txt)
            r.font.name = "Arial"
            r.font.size = Pt(12)
            r.bold   = bold
            r.italic = italic

        if _is_header(s):
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
            r = p.add_run(s)
            r.font.name = "Arial"
            r.font.size = Pt(13)
            r.bold = True
        else:
            p.paragraph_format.space_after = Pt(5)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            m_resp = _RE_RESPOSTA.match(s)
            m_ques = _RE_QUESITO_N.match(s)
            if m_resp:
                _run(m_resp.group(1), bold=True)
                resto = m_resp.group(2).strip()
                if resto:
                    _run(" " + resto, italic=True)
            elif m_ques:
                _run(m_ques.group(1), bold=True)
                _run(m_ques.group(2))
            else:
                p.paragraph_format.first_line_indent = Cm(1.25)
                _run(s)

    # ── Anexar PDFs de avaliação ──────────────────────────────────────────────
    if avaliacoes_paths:
        doc.add_page_break()
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(18)
        ph.paragraph_format.space_after  = Pt(6)
        rh = ph.add_run("ANEXOS — AVALIAÇÕES TÉCNICAS")
        rh.bold = True
        rh.font.name = "Arial"
        rh.font.size = Pt(13)

        for aval_path in avaliacoes_paths:
            aval_path = Path(aval_path)
            if aval_path.suffix.lower() != ".pdf":
                continue
            pt = doc.add_paragraph()
            pt.paragraph_format.space_before = Pt(10)
            pt.paragraph_format.space_after  = Pt(4)
            rt = pt.add_run(aval_path.stem)
            rt.bold = True
            rt.font.name = "Arial"
            rt.font.size = Pt(11)

            pages = pdf_para_imagens(str(aval_path), dpi=150)
            for img in pages:
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                buf.seek(0)
                pi = doc.add_paragraph()
                pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pi.paragraph_format.space_before = Pt(4)
                pi.paragraph_format.space_after  = Pt(4)
                pi.add_run().add_picture(buf, width=Cm(15.0))

    doc.save(caminho)


def construir_nome_arquivo(reclamante: str, reclamada: str,
                            insalubr: bool, periculos: bool, funcao: str) -> str:
    """Retorna: Reclamante x Reclamada = Tipo = Função"""
    rec, red, fun = reclamante.strip(), reclamada.strip(), funcao.strip()
    if rec and red:
        base = f"{rec} x {red}"
    elif rec:
        base = rec
    elif red:
        base = red
    else:
        return "LAUDO_PERICIAL"
    parts = [base]
    if insalubr:
        parts.append("Insalubr.")
    if periculos:
        parts.append("Periculos.")
    if fun:
        parts.append(fun)
    name = re.sub(r'[<>:"/\\|?*\n\r\t]', '', " = ".join(parts)).strip()
    return name[:180] or "LAUDO_PERICIAL"


def caminho_saida(pasta: str, prefixo: str) -> str:
    """Retorna caminho único para o .docx, respeitando MAX_PATH do Windows."""
    folder_len = len(str(Path(pasta)))
    max_nome = max(20, 259 - folder_len - 6)
    prefixo = prefixo[:max_nome]
    p = Path(pasta) / f"{prefixo}.docx"
    n = 2
    while p.exists():
        sfx = f"_{n}"
        p = Path(pasta) / f"{prefixo[:max_nome - len(sfx)]}{sfx}.docx"
        n += 1
    return str(p)

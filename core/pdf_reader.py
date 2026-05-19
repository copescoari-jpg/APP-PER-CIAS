"""
Leitura de PDF usando PyMuPDF (fitz) — única biblioteca PDF permitida.
"""

from pathlib import Path


def extrair_texto_pdf(path: str | Path, max_chars: int = 0) -> str:
    """Extrai texto de um PDF usando PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(str(path))
        partes = []
        for page in doc:
            t = page.get_text()
            if t.strip():
                partes.append(t)
        doc.close()
        texto = "\n".join(partes)
        return texto[:max_chars] if max_chars else texto
    except Exception as e:
        return f"[Erro ao ler PDF: {e}]"


def pdf_para_imagens(path: str | Path, dpi: int = 120, max_paginas: int = 4) -> list:
    """
    Converte páginas de um PDF em PIL.Image (RGB).
    Retorna lista vazia se fitz não estiver instalado.
    """
    try:
        import fitz
        from PIL import Image
        doc = fitz.open(str(path))
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        imgs = []
        for page in list(doc)[:max_paginas]:
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            imgs.append(img)
        doc.close()
        return imgs
    except Exception:
        return []


def extrair_texto_docx(path: str | Path) -> str:
    """Extrai texto de .docx preservando estrutura de parágrafos e tabelas."""
    try:
        from docx import Document
        doc = Document(str(path))
        linhas = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                linhas.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    linhas.append(" | ".join(cells))
        return "\n".join(linhas)
    except Exception as e:
        return f"[Erro ao ler DOCX: {e}]"


def extrair_texto(path: str | Path) -> str:
    """Roteador: extrai texto de .docx ou .pdf."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".docx":
        return extrair_texto_docx(p)
    if ext == ".pdf":
        return extrair_texto_pdf(p)
    return f"[Formato não suportado: {ext}]"
